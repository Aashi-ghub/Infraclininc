#!/usr/bin/env python3
"""
Generate Backend API Inventory DOCX Document

This script analyzes the backend codebase and generates a comprehensive
inventory document of all APIs, handlers, and storage dependencies.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json

# Define all API routes from serverless.ts analysis
api_inventory = [
    # Authentication APIs
    {
        "route": "/auth/login",
        "method": "POST",
        "handler": "auth.login",
        "file": "src/handlers/auth.ts",
        "domain": "Auth",
        "type": "AUTH",
        "storage": "NEITHER",
        "criticality": "CRITICAL",
        "notes": "Uses JSON file storage (temporary, will migrate to Cognito)"
    },
    {
        "route": "/auth/register",
        "method": "POST",
        "handler": "auth.register",
        "file": "src/handlers/auth.ts",
        "domain": "Auth",
        "type": "AUTH",
        "storage": "NEITHER",
        "criticality": "IMPORTANT",
        "notes": "Uses JSON file storage (temporary, will migrate to Cognito)"
    },
    {
        "route": "/auth/me",
        "method": "GET",
        "handler": "auth.me",
        "file": "src/handlers/auth.ts",
        "domain": "Auth",
        "type": "AUTH",
        "storage": "NEITHER",
        "criticality": "CRITICAL",
        "notes": "Uses JSON file storage (temporary, will migrate to Cognito)"
    },
    
    # Project APIs
    {
        "route": "/projects",
        "method": "GET",
        "handler": "projects.listProjects",
        "file": "src/handlers/projects.ts",
        "domain": "Project",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/projects/{project_id}",
        "method": "GET",
        "handler": "projects.getProject",
        "file": "src/handlers/projects.ts",
        "domain": "Project",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/projects",
        "method": "POST",
        "handler": "createProject.handler",
        "file": "src/handlers/createProject.ts",
        "domain": "Project",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Structure APIs
    {
        "route": "/structures",
        "method": "GET",
        "handler": "listStructures.handler",
        "file": "src/handlers/listStructures.ts",
        "domain": "Structure",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/structures",
        "method": "POST",
        "handler": "createStructure.handler",
        "file": "src/handlers/createStructure.ts",
        "domain": "Structure",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/structures/{structure_id}",
        "method": "GET",
        "handler": "getStructureById.handler",
        "file": "src/handlers/getStructureById.ts",
        "domain": "Structure",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Substructure APIs
    {
        "route": "/substructures",
        "method": "GET",
        "handler": "listSubstructures.handler",
        "file": "src/handlers/listSubstructures.ts",
        "domain": "SubStructure",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/substructures",
        "method": "POST",
        "handler": "createSubstructure.handler",
        "file": "src/handlers/createSubstructure.ts",
        "domain": "SubStructure",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/substructures/{substructure_id}",
        "method": "GET",
        "handler": "getSubstructureById.handler",
        "file": "src/handlers/getSubstructureById.ts",
        "domain": "SubStructure",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Geological Log APIs (Legacy)
    {
        "route": "/geological-log",
        "method": "POST",
        "handler": "createGeologicalLog.handler",
        "file": "src/handlers/createGeologicalLog.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    {
        "route": "/geological-log/{borelog_id}",
        "method": "GET",
        "handler": "getGeologicalLogById.handler",
        "file": "src/handlers/getGeologicalLogById.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    {
        "route": "/geological-log/{borelog_id}",
        "method": "PUT",
        "handler": "updateGeologicalLog.handler",
        "file": "src/handlers/updateGeologicalLog.ts",
        "domain": "Borelog",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    {
        "route": "/geological-log/{borelog_id}",
        "method": "DELETE",
        "handler": "deleteGeologicalLog.handler",
        "file": "src/handlers/deleteGeologicalLog.ts",
        "domain": "Borelog",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    {
        "route": "/geological-log",
        "method": "GET",
        "handler": "listGeologicalLogs.handler",
        "file": "src/handlers/listGeologicalLogs.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    {
        "route": "/geological-log/project-name/{project_name}",
        "method": "GET",
        "handler": "getGeologicalLogsByProjectName.handler",
        "file": "src/handlers/getGeologicalLogsByProjectName.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    {
        "route": "/geological-log/project-name/{project_name}/with-substructures",
        "method": "GET",
        "handler": "getGeologicalLogsByProjectNameWithSubstructures.handler",
        "file": "src/handlers/getGeologicalLogsByProjectNameWithSubstructures.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    {
        "route": "/geological-log/{borelog_id}/substructure",
        "method": "PUT",
        "handler": "updateSubstructureAssignment.handler",
        "file": "src/handlers/updateSubstructureAssignment.ts",
        "domain": "Borelog",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "Legacy endpoint - DB only"
    },
    
    # Borelog APIs
    {
        "route": "/borelog",
        "method": "POST",
        "handler": "createBorelog.handler",
        "file": "src/handlers/createBorelog.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/{borelog_id}",
        "method": "GET",
        "handler": "getBorelogBasicInfo.handler",
        "file": "src/handlers/getBorelogBasicInfo.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/{borelog_id}",
        "method": "DELETE",
        "handler": "deleteBorelog.handler",
        "file": "src/handlers/deleteBorelog.ts",
        "domain": "Borelog",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/substructure/{substructure_id}",
        "method": "GET",
        "handler": "getBorelogBySubstructureId.handler",
        "file": "src/handlers/getBorelogBySubstructureId.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/version",
        "method": "POST",
        "handler": "createBorelogVersion.handler",
        "file": "src/handlers/createBorelogVersion.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/upload-csv",
        "method": "POST",
        "handler": "uploadBorelogCSV.handler",
        "file": "src/handlers/uploadBorelogCSV.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "BOTH",
        "criticality": "CRITICAL",
        "notes": "Uploads CSV to S3, stores metadata in DB"
    },
    {
        "route": "/api/borelog/upload-csv",
        "method": "POST",
        "handler": "uploadBoreholeCsv.handler",
        "file": "src/handlers/uploadBoreholeCsv.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "BOTH",
        "criticality": "IMPORTANT",
        "notes": "Alternative CSV upload endpoint - uploads to S3, stores metadata in DB"
    },
    {
        "route": "/borelog/{borelog_id}/approve",
        "method": "POST",
        "handler": "approveBorelog.handler",
        "file": "src/handlers/approveBorelog.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/submit",
        "method": "POST",
        "handler": "borelogSubmission.submitBorelog",
        "file": "src/handlers/borelogSubmission.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/submissions/{projectId}/{boreholeId}",
        "method": "GET",
        "handler": "borelogSubmission.getBorelogSubmissions",
        "file": "src/handlers/borelogSubmission.ts",
        "domain": "Approval",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog/submission/{submissionId}",
        "method": "GET",
        "handler": "borelogSubmission.getBorelogSubmission",
        "file": "src/handlers/borelogSubmission.ts",
        "domain": "Approval",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelogs",
        "method": "GET",
        "handler": "listGeologicalLogs.handler",
        "file": "src/handlers/listGeologicalLogs.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "Alias for /geological-log - DB only"
    },
    {
        "route": "/projects/{project_id}/borelogs",
        "method": "GET",
        "handler": "getBorelogsByProject.handler",
        "file": "src/handlers/getBorelogsByProject.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Borelog Details APIs
    {
        "route": "/borelog-details",
        "method": "POST",
        "handler": "createBorelogDetails.handler",
        "file": "src/handlers/createBorelogDetails.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog-details/{borelog_id}",
        "method": "GET",
        "handler": "getBorelogDetailsByBorelogId.handler",
        "file": "src/handlers/getBorelogDetailsByBorelogId.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Stratum Data APIs
    {
        "route": "/stratum-data",
        "method": "POST",
        "handler": "saveStratumData.handler",
        "file": "src/handlers/saveStratumData.ts",
        "domain": "Stratum",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/stratum-data",
        "method": "GET",
        "handler": "getStratumData.handler",
        "file": "src/handlers/getStratumData.ts",
        "domain": "Stratum",
        "type": "READ",
        "storage": "S3",
        "criticality": "CRITICAL",
        "notes": "MIGRATED TO S3 - Reads from S3 Parquet storage"
    },
    
    # Borelog Form Data API
    {
        "route": "/borelog-form-data",
        "method": "GET",
        "handler": "getBorelogFormData.handler",
        "file": "src/handlers/getBorelogFormData.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Borelog Images APIs
    {
        "route": "/borelog-images",
        "method": "POST",
        "handler": "borelogImages.uploadImage",
        "file": "src/handlers/borelogImages.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "BOTH",
        "criticality": "IMPORTANT",
        "notes": "Uploads images to S3, stores metadata in DB"
    },
    {
        "route": "/borelog-images/{borelog_id}",
        "method": "GET",
        "handler": "borelogImages.getImages",
        "file": "src/handlers/borelogImages.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "Reads metadata from DB, images served from S3"
    },
    {
        "route": "/borelog-images/{image_id}",
        "method": "DELETE",
        "handler": "borelogImages.deleteImage",
        "file": "src/handlers/borelogImages.ts",
        "domain": "Borelog",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "Deletes metadata from DB, image remains in S3"
    },
    
    # Borelog Assignment APIs
    {
        "route": "/borelog-assignments",
        "method": "POST",
        "handler": "borelogAssignments.createAssignment",
        "file": "src/handlers/borelogAssignments.ts",
        "domain": "Assignment",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog-assignments/{assignmentId}",
        "method": "PUT",
        "handler": "borelogAssignments.updateAssignment",
        "file": "src/handlers/borelogAssignments.ts",
        "domain": "Assignment",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog-assignments/borelog/{borelogId}",
        "method": "GET",
        "handler": "borelogAssignments.getAssignmentsByBorelogId",
        "file": "src/handlers/borelogAssignments.ts",
        "domain": "Assignment",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog-assignments/structure/{structureId}",
        "method": "GET",
        "handler": "borelogAssignments.getAssignmentsByStructureId",
        "file": "src/handlers/borelogAssignments.ts",
        "domain": "Assignment",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog-assignments/site-engineer/{siteEngineerId}",
        "method": "GET",
        "handler": "borelogAssignments.getAssignmentsBySiteEngineer",
        "file": "src/handlers/borelogAssignments.ts",
        "domain": "Assignment",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog-assignments/active",
        "method": "GET",
        "handler": "borelogAssignments.getActiveAssignments",
        "file": "src/handlers/borelogAssignments.ts",
        "domain": "Assignment",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/borelog-assignments/{assignmentId}",
        "method": "DELETE",
        "handler": "borelogAssignments.deleteAssignment",
        "file": "src/handlers/borelogAssignments.ts",
        "domain": "Assignment",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # User Assignment APIs
    {
        "route": "/assignments",
        "method": "POST",
        "handler": "assignUsers.handler",
        "file": "src/handlers/assignUsers.ts",
        "domain": "Assignment",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # User Management APIs
    {
        "route": "/users",
        "method": "GET",
        "handler": "users.listUsers",
        "file": "src/handlers/users.ts",
        "domain": "Auth",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/users/{user_id}",
        "method": "GET",
        "handler": "users.getUserById",
        "file": "src/handlers/users.ts",
        "domain": "Auth",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/users/lab-engineers",
        "method": "GET",
        "handler": "users.getLabEngineers",
        "file": "src/handlers/users.ts",
        "domain": "Auth",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    
    # Lab Test APIs
    {
        "route": "/lab-tests",
        "method": "POST",
        "handler": "labTests.createLabTest",
        "file": "src/handlers/labTests.ts",
        "domain": "Lab Test",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-tests",
        "method": "GET",
        "handler": "labTests.listLabTests",
        "file": "src/handlers/labTests.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Lab Request APIs
    {
        "route": "/lab-requests",
        "method": "POST",
        "handler": "labRequests.createLabRequest",
        "file": "src/handlers/labRequests.ts",
        "domain": "Lab Test",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-requests",
        "method": "GET",
        "handler": "labRequests.listLabRequests",
        "file": "src/handlers/labRequests.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-requests/{id}",
        "method": "GET",
        "handler": "labRequests.getLabRequestById",
        "file": "src/handlers/labRequests.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-requests/{id}",
        "method": "PUT",
        "handler": "labRequests.updateLabRequest",
        "file": "src/handlers/labRequests.ts",
        "domain": "Lab Test",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-requests/{id}",
        "method": "DELETE",
        "handler": "labRequests.deleteLabRequest",
        "file": "src/handlers/labRequests.ts",
        "domain": "Lab Test",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-requests/final-borelogs",
        "method": "GET",
        "handler": "labRequests.getFinalBorelogs",
        "file": "src/handlers/labRequests.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    
    # Unified Lab Reports APIs
    {
        "route": "/unified-lab-reports",
        "method": "POST",
        "handler": "unifiedLabReports.createUnifiedLabReport",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Report",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports/{reportId}",
        "method": "GET",
        "handler": "unifiedLabReports.getUnifiedLabReport",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Report",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports/{reportId}",
        "method": "PUT",
        "handler": "unifiedLabReports.updateUnifiedLabReport",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Report",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports",
        "method": "GET",
        "handler": "unifiedLabReports.getUnifiedLabReports",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Report",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports/{reportId}",
        "method": "DELETE",
        "handler": "unifiedLabReports.deleteUnifiedLabReport",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Report",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports/{reportId}/approve",
        "method": "POST",
        "handler": "unifiedLabReports.approveUnifiedLabReport",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports/{reportId}/reject",
        "method": "POST",
        "handler": "unifiedLabReports.rejectUnifiedLabReport",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports/{reportId}/submit",
        "method": "POST",
        "handler": "unifiedLabReports.submitUnifiedLabReport",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/unified-lab-reports/upload-csv",
        "method": "POST",
        "handler": "uploadUnifiedLabReportCSV.handler",
        "file": "src/handlers/uploadUnifiedLabReportCSV.ts",
        "domain": "Report",
        "type": "CREATE",
        "storage": "BOTH",
        "criticality": "CRITICAL",
        "notes": "Uploads CSV to S3, stores metadata in DB"
    },
    {
        "route": "/lab-reports",
        "method": "GET",
        "handler": "unifiedLabReports.getUnifiedLabReports",
        "file": "src/handlers/unifiedLabReports.ts",
        "domain": "Report",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "Alias for /unified-lab-reports - DB only"
    },
    
    # Lab Report Version Control APIs
    {
        "route": "/lab-reports/draft",
        "method": "POST",
        "handler": "labReportVersionControl.saveLabReportDraft",
        "file": "src/handlers/labReportVersionControl.ts",
        "domain": "Report",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-reports/submit",
        "method": "POST",
        "handler": "labReportVersionControl.submitLabReportForReview",
        "file": "src/handlers/labReportVersionControl.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-reports/{report_id}/review",
        "method": "POST",
        "handler": "labReportVersionControl.reviewLabReport",
        "file": "src/handlers/labReportVersionControl.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-reports/{report_id}/versions",
        "method": "GET",
        "handler": "labReportVersionControl.getLabReportVersionHistory",
        "file": "src/handlers/labReportVersionControl.ts",
        "domain": "Report",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/lab-reports/{report_id}/version/{version_no}",
        "method": "GET",
        "handler": "labReportVersionControl.getLabReportVersion",
        "file": "src/handlers/labReportVersionControl.ts",
        "domain": "Report",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    
    # Soil Test Samples APIs
    {
        "route": "/unified-lab-reports/{reportId}/soil-samples",
        "method": "GET",
        "handler": "soilTestSamples.getSoilTestSamples",
        "file": "src/handlers/soilTestSamples.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/soil-test-samples/{sampleId}",
        "method": "GET",
        "handler": "soilTestSamples.getSoilTestSample",
        "file": "src/handlers/soilTestSamples.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/soil-test-samples/{sampleId}",
        "method": "PUT",
        "handler": "soilTestSamples.updateSoilTestSample",
        "file": "src/handlers/soilTestSamples.ts",
        "domain": "Lab Test",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/soil-test-samples/{sampleId}",
        "method": "DELETE",
        "handler": "soilTestSamples.deleteSoilTestSample",
        "file": "src/handlers/soilTestSamples.ts",
        "domain": "Lab Test",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    
    # Rock Test Samples APIs
    {
        "route": "/unified-lab-reports/{reportId}/rock-samples",
        "method": "GET",
        "handler": "rockTestSamples.getRockTestSamples",
        "file": "src/handlers/rockTestSamples.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/rock-test-samples/{sampleId}",
        "method": "GET",
        "handler": "rockTestSamples.getRockTestSample",
        "file": "src/handlers/rockTestSamples.ts",
        "domain": "Lab Test",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/rock-test-samples/{sampleId}",
        "method": "PUT",
        "handler": "rockTestSamples.updateRockTestSample",
        "file": "src/handlers/rockTestSamples.ts",
        "domain": "Lab Test",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/rock-test-samples/{sampleId}",
        "method": "DELETE",
        "handler": "rockTestSamples.deleteRockTestSample",
        "file": "src/handlers/rockTestSamples.ts",
        "domain": "Lab Test",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    
    # Workflow APIs
    {
        "route": "/workflow/{borelog_id}/submit",
        "method": "POST",
        "handler": "workflowActions.submitForReview",
        "file": "src/handlers/workflowActions.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/{borelog_id}/review",
        "method": "POST",
        "handler": "workflowActions.reviewBorelog",
        "file": "src/handlers/workflowActions.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/lab-assignments",
        "method": "POST",
        "handler": "workflowActions.assignLabTests",
        "file": "src/handlers/workflowActions.ts",
        "domain": "Assignment",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/lab-results",
        "method": "POST",
        "handler": "workflowActions.submitLabTestResults",
        "file": "src/handlers/workflowActions.ts",
        "domain": "Lab Test",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/{borelog_id}/status",
        "method": "GET",
        "handler": "workflowActions.getWorkflowStatus",
        "file": "src/handlers/workflowActions.ts",
        "domain": "Approval",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/pending-reviews",
        "method": "GET",
        "handler": "workflowDashboard.getPendingReviews",
        "file": "src/handlers/workflowDashboard.ts",
        "domain": "Approval",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/lab-assignments",
        "method": "GET",
        "handler": "workflowDashboard.getLabAssignments",
        "file": "src/handlers/workflowDashboard.ts",
        "domain": "Assignment",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/statistics",
        "method": "GET",
        "handler": "workflowDashboard.getWorkflowStatistics",
        "file": "src/handlers/workflowDashboard.ts",
        "domain": "Approval",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/workflow/submitted-borelogs",
        "method": "GET",
        "handler": "workflowDashboard.getSubmittedBorelogs",
        "file": "src/handlers/workflowDashboard.ts",
        "domain": "Approval",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Anomaly APIs
    {
        "route": "/anomalies",
        "method": "GET",
        "handler": "anomalies.listAnomalies",
        "file": "src/handlers/anomalies.ts",
        "domain": "Approval",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/anomalies",
        "method": "POST",
        "handler": "anomalies.createAnomaly",
        "file": "src/handlers/anomalies.ts",
        "domain": "Approval",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/anomalies/{anomaly_id}",
        "method": "PATCH",
        "handler": "anomalies.updateAnomaly",
        "file": "src/handlers/anomalies.ts",
        "domain": "Approval",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    
    # Contact APIs
    {
        "route": "/contacts",
        "method": "POST",
        "handler": "contacts.createContactHandler",
        "file": "src/handlers/contacts.ts",
        "domain": "Project",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/contacts",
        "method": "GET",
        "handler": "contacts.listContactsHandler",
        "file": "src/handlers/contacts.ts",
        "domain": "Project",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/contacts/{contact_id}",
        "method": "GET",
        "handler": "contacts.getContactByIdHandler",
        "file": "src/handlers/contacts.ts",
        "domain": "Project",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/contacts/organisation/{organisation_id}",
        "method": "GET",
        "handler": "contacts.getContactsByOrganisationHandler",
        "file": "src/handlers/contacts.ts",
        "domain": "Project",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/contacts/{contact_id}",
        "method": "PUT",
        "handler": "contacts.updateContactHandler",
        "file": "src/handlers/contacts.ts",
        "domain": "Project",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/contacts/{contact_id}",
        "method": "DELETE",
        "handler": "contacts.deleteContactHandler",
        "file": "src/handlers/contacts.ts",
        "domain": "Project",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "OPTIONAL",
        "notes": "DB only - needs migration to S3"
    },
    
    # Borehole APIs
    {
        "route": "/boreholes",
        "method": "GET",
        "handler": "boreholes.listBoreholes",
        "file": "src/handlers/boreholes.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/boreholes/{boreholeId}",
        "method": "GET",
        "handler": "boreholes.getBoreholeById",
        "file": "src/handlers/boreholes.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/boreholes/project/{projectId}",
        "method": "GET",
        "handler": "boreholes.getBoreholesByProject",
        "file": "src/handlers/boreholes.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/boreholes/project/{projectId}/structure/{structureId}",
        "method": "GET",
        "handler": "boreholes.getBoreholesByProjectAndStructure",
        "file": "src/handlers/boreholes.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/boreholes",
        "method": "POST",
        "handler": "boreholes.createBorehole",
        "file": "src/handlers/boreholes.ts",
        "domain": "Borelog",
        "type": "CREATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/boreholes/{boreholeId}",
        "method": "PUT",
        "handler": "boreholes.updateBorehole",
        "file": "src/handlers/boreholes.ts",
        "domain": "Borelog",
        "type": "UPDATE",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/boreholes/{boreholeId}",
        "method": "DELETE",
        "handler": "boreholes.deleteBorehole",
        "file": "src/handlers/boreholes.ts",
        "domain": "Borelog",
        "type": "DELETE",
        "storage": "DATABASE",
        "criticality": "IMPORTANT",
        "notes": "DB only - needs migration to S3"
    },
    
    # Pending CSV Upload APIs
    {
        "route": "/pending-csv-uploads",
        "method": "GET",
        "handler": "listPendingCSVUploads.handler",
        "file": "src/handlers/listPendingCSVUploads.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/pending-csv-uploads/{upload_id}",
        "method": "GET",
        "handler": "getPendingCSVUpload.handler",
        "file": "src/handlers/getPendingCSVUpload.ts",
        "domain": "Borelog",
        "type": "READ",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
    {
        "route": "/pending-csv-uploads/{upload_id}/approve",
        "method": "POST",
        "handler": "approvePendingCSVUpload.handler",
        "file": "src/handlers/approvePendingCSVUpload.ts",
        "domain": "Approval",
        "type": "APPROVAL",
        "storage": "DATABASE",
        "criticality": "CRITICAL",
        "notes": "DB only - needs migration to S3"
    },
]

def create_document():
    """Create the DOCX document with API inventory"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Backend API Inventory and Migration Map', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata paragraph
    meta = doc.add_paragraph()
    meta.add_run('Generated: ').bold = True
    from datetime import datetime
    meta.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # 1. Backend Overview
    doc.add_heading('1. Backend Overview', 1)
    overview = doc.add_paragraph()
    overview.add_run('Architecture: ').bold = True
    overview.add_run('Serverless backend built with AWS Lambda, API Gateway, and TypeScript. Uses PostgreSQL database (currently disabled) and S3 + Parquet storage (active).')
    doc.add_paragraph()
    overview2 = doc.add_paragraph()
    overview2.add_run('Current Storage State: ').bold = True
    overview2.add_run('Database (PostgreSQL) is disabled (DB_ENABLED=false). S3 storage with Parquet format is active for stratum data. Most APIs still depend on database and need migration.')
    doc.add_paragraph()
    
    # 2. API Inventory Table
    doc.add_heading('2. API Inventory Table', 1)
    doc.add_paragraph('Complete inventory of all backend APIs with their storage dependencies and classifications.')
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Route', 'Method', 'Handler', 'File', 'Domain', 'Type', 'Storage', 'Criticality']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add data rows
    for api in api_inventory:
        row_cells = table.add_row().cells
        row_cells[0].text = api['route']
        row_cells[1].text = api['method']
        row_cells[2].text = api['handler']
        row_cells[3].text = api['file']
        row_cells[4].text = api['domain']
        row_cells[5].text = api['type']
        row_cells[6].text = api['storage']
        row_cells[7].text = api['criticality']
    
    doc.add_page_break()
    
    # 3. Creation Flow APIs
    doc.add_heading('3. Creation Flow APIs', 1)
    doc.add_paragraph('APIs involved in creating projects, structures, and borelogs.')
    doc.add_paragraph()
    
    creation_apis = [api for api in api_inventory if api['type'] == 'CREATE' and api['domain'] in ['Project', 'Structure', 'SubStructure', 'Borelog']]
    
    for api in creation_apis:
        p = doc.add_paragraph()
        p.add_run(f"{api['method']} {api['route']}").bold = True
        p.add_run(f" - {api['handler']} ({api['file']})")
        p.add_run(f"\nStorage: {api['storage']} | Criticality: {api['criticality']}")
        if api.get('notes'):
            p.add_run(f"\nNotes: {api['notes']}")
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 4. Data Entry APIs
    doc.add_heading('4. Data Entry APIs', 1)
    doc.add_paragraph('APIs for entering stratum data, borelog details, and lab tests.')
    doc.add_paragraph()
    
    data_entry_apis = [api for api in api_inventory if api['domain'] in ['Stratum', 'Lab Test'] and api['type'] in ['CREATE', 'UPDATE']]
    
    for api in data_entry_apis:
        p = doc.add_paragraph()
        p.add_run(f"{api['method']} {api['route']}").bold = True
        p.add_run(f" - {api['handler']} ({api['file']})")
        p.add_run(f"\nStorage: {api['storage']} | Criticality: {api['criticality']}")
        if api.get('notes'):
            p.add_run(f"\nNotes: {api['notes']}")
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 5. Read / View APIs
    doc.add_heading('5. Read / View APIs', 1)
    doc.add_paragraph('APIs used for dashboard, borelog view, and reports.')
    doc.add_paragraph()
    
    read_apis = [api for api in api_inventory if api['type'] == 'READ']
    
    # Group by domain
    domains = {}
    for api in read_apis:
        domain = api['domain']
        if domain not in domains:
            domains[domain] = []
        domains[domain].append(api)
    
    for domain, apis in sorted(domains.items()):
        doc.add_heading(f'{domain} Domain', 2)
        for api in apis:
            p = doc.add_paragraph()
            p.add_run(f"{api['method']} {api['route']}").bold = True
            p.add_run(f" - {api['handler']}")
            p.add_run(f"\nStorage: {api['storage']} | Criticality: {api['criticality']}")
            if api.get('notes'):
                p.add_run(f"\nNotes: {api['notes']}")
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # 6. Approval & Workflow APIs
    doc.add_heading('6. Approval & Workflow APIs', 1)
    doc.add_paragraph('APIs related to approval, assignment, and status changes.')
    doc.add_paragraph()
    
    approval_apis = [api for api in api_inventory if api['type'] == 'APPROVAL' or api['domain'] == 'Approval' or api['domain'] == 'Assignment']
    
    for api in approval_apis:
        p = doc.add_paragraph()
        p.add_run(f"{api['method']} {api['route']}").bold = True
        p.add_run(f" - {api['handler']} ({api['file']})")
        p.add_run(f"\nStorage: {api['storage']} | Criticality: {api['criticality']}")
        if api.get('notes'):
            p.add_run(f"\nNotes: {api['notes']}")
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 7. Summary & Migration Notes
    doc.add_heading('7. Summary & Migration Notes', 1)
    
    # Count APIs by storage type
    db_only = [api for api in api_inventory if api['storage'] == 'DATABASE']
    s3_only = [api for api in api_inventory if api['storage'] == 'S3']
    both = [api for api in api_inventory if api['storage'] == 'BOTH']
    neither = [api for api in api_inventory if api['storage'] == 'NEITHER']
    
    doc.add_heading('7.1 Storage Dependency Summary', 2)
    summary_p = doc.add_paragraph()
    summary_p.add_run(f'Total APIs: {len(api_inventory)}\n').bold = True
    summary_p.add_run(f'• Database Only: {len(db_only)} APIs\n')
    summary_p.add_run(f'• S3 Only: {len(s3_only)} APIs\n')
    summary_p.add_run(f'• Both DB and S3: {len(both)} APIs\n')
    summary_p.add_run(f'• Neither (Pure Logic): {len(neither)} APIs\n')
    
    doc.add_paragraph()
    
    doc.add_heading('7.2 APIs Still Dependent on Database', 2)
    doc.add_paragraph(f'{len(db_only)} APIs still require database access and need migration to S3:')
    doc.add_paragraph()
    
    for api in db_only[:20]:  # Show first 20
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{api['method']} {api['route']}").bold = True
        p.add_run(f" ({api['domain']} - {api['type']})")
    
    if len(db_only) > 20:
        doc.add_paragraph(f'... and {len(db_only) - 20} more APIs')
    
    doc.add_paragraph()
    
    doc.add_heading('7.3 APIs Already Using S3', 2)
    doc.add_paragraph('APIs that are already using S3 storage:')
    doc.add_paragraph()
    
    s3_apis = s3_only + both
    for api in s3_apis:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{api['method']} {api['route']}").bold = True
        p.add_run(f" - {api['storage']}")
    
    doc.add_paragraph()
    
    doc.add_heading('7.4 Suggested Migration Order', 2)
    doc.add_paragraph('High-level migration strategy (no code changes):')
    doc.add_paragraph()
    
    migration_steps = [
        ('Phase 1: Core Data', [
            'Stratum data (already migrated)',
            'Borelog basic info',
            'Borelog details',
            'Borelog versions'
        ]),
        ('Phase 2: Project Structure', [
            'Projects',
            'Structures',
            'Substructures',
            'Boreholes'
        ]),
        ('Phase 3: Workflow & Assignments', [
            'Borelog assignments',
            'User assignments',
            'Workflow status',
            'Approval workflows'
        ]),
        ('Phase 4: Lab Data', [
            'Lab tests',
            'Lab requests',
            'Lab reports',
            'Test samples (soil & rock)'
        ]),
        ('Phase 5: Supporting Data', [
            'Users (consider Cognito migration)',
            'Contacts',
            'Anomalies',
            'Images metadata'
        ])
    ]
    
    for phase_name, items in migration_steps:
        doc.add_heading(phase_name, 3)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)
        doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph('Note: This is a high-level migration plan. Actual implementation should be done incrementally with thorough testing at each phase.')
    
    return doc

if __name__ == '__main__':
    print("Generating Backend API Inventory DOCX document...")
    doc = create_document()
    output_path = 'Backend_API_Inventory_and_Migration_Map.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Total APIs documented: {len(api_inventory)}")









